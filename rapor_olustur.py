# -*- coding: utf-8 -*-
"""NSosyal Duygu Katmanı - TEKNOFEST Teknik Rapor üretici.
Resmi şablona (NSosyal_Inovasyon_2026_-_Proje_Teknik_Raporu_1_u6IVb) uyar:
Arial 12pt gövde, Arial Black 14pt başlık, 1.15 satır aralığı, iki yana yaslı,
2.5cm kenar boşluğu, kapak+içindekiler+kaynakça ayrı sayfa, <=30 sayfa."""
import docx
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOSYA = "NSosyal_Teknik_Rapor.docx"

doc = Document()

# ---------- Sayfa / stil kurulumu ----------
for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def _font_ayarla(style, ad, kalin=None):
    style.font.name = ad
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:ascii'), ad)
    rFonts.set(qn('w:hAnsi'), ad)
    rFonts.set(qn('w:cs'), ad)
    if kalin is not None:
        style.font.bold = kalin


normal = doc.styles['Normal']
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(6)
_font_ayarla(normal, 'Arial')

h1 = doc.styles['Heading 1']
h1.font.size = Pt(14)
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.space_before = Pt(20)
h1.paragraph_format.space_after = Pt(10)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1.paragraph_format.line_spacing = 1.15
h1.paragraph_format.keep_with_next = True
_font_ayarla(h1, 'Arial Black', kalin=True)

# ---------- Yardımcı fonksiyonlar ----------

def baslik(metin):
    doc.add_heading(metin, level=1)


def altbaslik(metin):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(metin)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(12.5)
    rpr = r._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rpr.append(rFonts)


def govde(metin):
    doc.add_paragraph(metin)


def madde(satirlar):
    for s in satirlar:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(s)
        r.font.name = 'Arial'
        r.font.size = Pt(12)


def sayfa_sonu():
    doc.add_page_break()


def notv(metin):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(metin)
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.name = 'Arial'


def tablo(basliklar, satirlar):
    t = doc.add_table(rows=1, cols=len(basliklar))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(basliklar):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
    for satir in satirlar:
        cells = t.add_row().cells
        for i, deger in enumerate(satir):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(deger))
            r.font.name = 'Arial'
            r.font.size = Pt(10.5)
    doc.add_paragraph()
    return t


def toc_ekle():
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-1" \\h \\z \\u'
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'separate')
    metin = OxmlElement('w:t')
    metin.text = "İçindekiler listesi görünmüyorsa: sağ tık > Alanı Güncelle (veya F9)."
    fld3 = OxmlElement('w:fldChar')
    fld3.set(qn('w:fldCharType'), 'end')
    r_el = run._r
    r_el.append(fld1)
    r_el.append(instr)
    r_el.append(fld2)
    r_el.append(metin)
    r_el.append(fld3)


# ============================================================ KAPAK
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(140)
r = p.add_run("NSosyal Duygu Katmanı")
r.bold = True
r.font.size = Pt(26)
r.font.name = 'Arial Black'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Duygu-Duyarlı, Açıklanabilir ve Koruyucu Bir Sıralama Katmanı")
r.font.size = Pt(15)
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("TEKNOFEST NSosyal İnovasyon Yarışması 2026")
r.font.size = Pt(13)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Tema: Sosyal Yapay Zekâ")
r2.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(150)
r = p.add_run("Teknik Rapor")
r.font.size = Pt(13)
r.bold = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("24 Ağustos 2026")
r2.font.size = Pt(12)

sayfa_sonu()

# ============================================================ İÇİNDEKİLER
baslik("İÇİNDEKİLER")
toc_ekle()
sayfa_sonu()

# ============================================================ 1. PROJE ÖZETİ
baslik("PROJE ÖZETİ")

altbaslik("1.1. Proje Konusu ve Amacı")
govde(
    "Bu proje, NSosyal (T3 Vakfı ve Baykar Teknoloji iş birliğiyle geliştirilen, "
    "~700 bin'i aşkın aktif kullanıcısı bulunan, reklamsız bir yerli mikroblog platformu) "
    "için duygu-duyarlı, açıklanabilir ve koruyucu bir sıralama ve şeffaflık katmanı "
    "geliştirmeyi konu alır. Sosyal medya akışlarının çoğu \"en yüksek etkileşim\" hedefiyle "
    "optimize edilir; bu durum olumsuz ve kutuplaştırıcı içeriğin yapısal olarak "
    "ödüllendirilmesine yol açabilir. Bu soyut bir varsayım değildir: Facebook'un kendi iç "
    "araştırması, \"öfke\" tepkisi taşıyan gönderilerin çok daha fazla etkileşim aldığını "
    "ortaya koymuştur [13]; bağımsız, hakemli bir çalışma ise haber başlıklarındaki her ek "
    "negatif kelimenin tıklama oranını %2,3 artırdığını nedensel olarak göstermiştir [10]; "
    "kırılgan duygu durumlarının reklamcılık amacıyla hedeflenmesinin de belgelenmiş bir "
    "emsali bulunmaktadır [14] (ayrıntılar için bkz. Bölüm 2.1). Projenin amacı, kullanıcının ilgi "
    "alanından hiç sapmadan, yalnızca o an olumsuz bir davranışsal örüntüye (\"olası "
    "spiral\") girip girmediğini fark eden ve buna göre içerik dozunu nazikçe ayarlayan, "
    "kararlarının tamamı denetlenebilir/açıklanabilir bir sistem sunmaktır. Proje, birincil "
    "olarak Sosyal Yapay Zekâ temasına hitap eder; açıklanabilir sıralama motoru ve "
    "şeffaflık paneli üzerinden Kullanıcı Katılımı ve Arayüz/Kullanıcı Deneyimi (UI/UX) "
    "temasına, üretici paneli önerisiyle de İçerik Ekonomisi temasına organik olarak "
    "dokunur. Bu amaç, yarışmanın Bölüm 1'inde (Yarışma Amacı) sayılan temel hedeflerle "
    "doğrudan örtüşür: yapay zekâ destekli yeni nesil sosyal medya platformları "
    "geliştirilmesi ve güvenli, etik, şeffaf, kullanıcı mahremiyetini önceleyen çözümler "
    "hedefleri, projenin tam da adresi olduğu alanlardır. Aynı bölümde sosyal medyanın "
    "yaygınlaşmasının beraberinde getirdiği ihtiyaçlar arasında ayrıca açıkça sayılan "
    "kullanıcı refahının korunması ihtiyacı, bu projenin doğrudan çıkış noktasını "
    "oluşturmaktadır."
)

altbaslik("1.2. Proje Kapsamı ve Yöntemi")
govde(
    "Projenin kapsamı, NSosyal'in ana akışına eklenen üç bileşenden oluşur. Birincisi, "
    "gönderi metninden gerçek zamanlı duygu skoru çıkaran, bağımsız olarak doğrulanmış "
    "ve ince ayarlı bir Türkçe BERT modelidir. İkincisi, kullanıcının davranışsal "
    "sinyallerinden (durma süresi, tıklama, roket, yorum) olası bir \"spiral\" durumu "
    "ve olası bir psikolojik izlenim kategorisi tahmin eden, sentetik senaryo "
    "verisiyle eğitilmiş açıklanabilir sınıflandırıcılardır. Üçüncüsü, bu sinyalleri "
    "ilgi skoruyla birleştirip içeriği yeniden sıralayan iki katmanlı bir motordur. "
    "Yöntem olarak; önce mimari "
    "sözlük-tabanlı basit bir ilk prototiple (spike) uçtan uca doğrulanmış, ardından "
    "her bileşen gerçek modellerle (BERT, lojistik regresyon/SGD) değiştirilip bağımsız "
    "olarak ölçülmüştür. Fikir yalnızca tasarım düzeyinde kalmamış; FastAPI tabanlı bir "
    "backend, gerçek zamanlı bir web arayüzü, ve isteğe bağlı bir LLM destekli haftalık "
    "öz-farkındalık raporu içeren, uçtan uca çalışan bir prototiple desteklenmiştir "
    "(bkz. Bölüm 3-4). NSosyal veya T3 AI'a gerçek bir API erişimi bulunmamaktadır; "
    "sistem bağımsız bir prototip olarak çalışır ve \"önerilen entegrasyon "
    "konsepti\" olarak sunulur; bu durum, projenin dürüstlük ilkesi gereği açıkça belirtilir. "
    "Kapsam dışında kalan yönler de aynı netlikle çizilmelidir: proje içerik "
    "kaldırma/moderasyon kararı almaz, klinik bir teşhis veya tedavi önerisi sunmaz; "
    "yalnızca mevcut sıralama ve şeffaflık katmanını hedefler. Akademik yöntem olarak, "
    "tasarım kararları ilgili literatürle temellendirilmiş ve vendor/model iddiaları "
    "bağımsız veri setleriyle sınanmıştır (bkz. Bölüm 3.2); bu, projeyi yalnızca "
    "mühendislik değil, ölçülebilir bir doğrulama disipliniyle de temellendirir. Proje "
    "kendi kapsamı içinde tamamlanmış bir son ürün değil, bilinçli bir başlangıç "
    "noktasıdır: Bölüm 4.3'te ayrıntılandırılan kişiselleştirme, gizlilik öncelikli "
    "mimari ve perspektif genişletici öneriler gibi yönler için somut bir zemin hazırlar."
)

# ============================================================ 2. KATMA DEĞER VE YENİLİKÇİLİK
baslik("KATMA DEĞER VE YENİLİKÇİLİK")

altbaslik("2.1. Problem Tanımı ve Mevcut Çözümler")
govde(
    "Bölüm 1.1'de değinilen örüntünün arkasında belgelenmiş, kamuya mal olmuş kanıtlar "
    "vardır. Wall Street Journal'ın 2021'de yayımladığı \"The Facebook Files\" dizisi, "
    "şirketin \"öfke\" tepkisini 2017-2018'de algoritmada beğeniden 5 kat ağırlıklandırdığını, "
    "iç araştırmacıların bunun kutuplaştırıcı/yanlış bilgi içeren içeriği ödüllendirdiğini "
    "fark edip bu ağırlığı sonradan düşürdüğünü ortaya koymuştur [13]. Akademik tarafta, "
    "yaklaşık 105.000 haber başlığı varyasyonunun ~5,7 milyon tıklama üzerinden analiz "
    "edildiği randomize kontrollü bir çalışma, bu etkinin (%2,3'lük tıklama artışı) "
    "yalnızca bir korelasyon değil, nedensel bir ilişki olduğunu göstermiştir [10]; "
    "olumsuz haberlerin sosyal medyada daha fazla paylaşıldığı da "
    "ayrıca doğrulanmıştır [11][12]. Pasif/olumsuz içerik tüketiminin kaygı ve stresle "
    "ilişkisi de literatürde tutarlı biçimde destekleniyor: 141 çalışmayı kapsayan bir "
    "meta-analiz bu ilişkiyi orta güçte doğrularken [4], doomscrolling'e odaklanan bir "
    "başka çalışma bu davranışı kaygı, stres ve kontrol kaybı hissiyle "
    "ilişkilendirmektedir [5]. "
    "Kırılganlık durumunun ticari olarak istismar edilmesinin de gerçek bir emsali "
    "vardır: 2017'de The Australian gazetesi, Facebook'un iç bir belgesinin, gençlerin "
    "ne zaman \"değersiz\" veya \"güvensiz\" hissettiğini tespit edip bunu reklam "
    "hedeflemesi için bir fırsat olarak sunduğunu haberleştirmiştir [14]."
)
govde(
    "Piyasadaki mevcut çözümler, bu problemi doğrudan hedef almak yerine dolaylı ve "
    "genellikle isteğe bağlı (opt-in) önlemlerle sınırlı kalmaktadır. Instagram'ın "
    "2021'de tanıttığı \"Take a Break\" hatırlatıcıları, kullanıcıyı 10/20/30 dakikalık "
    "aralıklarla uygulamadan uzaklaşmaya davet eder [17]; TikTok da benzer şekilde "
    "\"Scheduled Breaks\" ve bir ekran-süresi paneli sunar [18]. Ancak bu araçlar iki "
    "ortak sınırlılık taşır. Birincisi, hangi içeriğin tüketildiğinden veya kullanıcının "
    "o an nasıl hissettiğinden tamamen bağımsızdır; yalnızca geçen süreye bakarlar. "
    "İkincisi, sıralama algoritmasının kendisine hiç dokunmazlar; yalnızca üzerine "
    "isteğe bağlı bir hatırlatma katmanı eklerler, kullanıcı görmezden gelebilir. Daha "
    "genel olarak, piyasadaki çözümler büyük ölçüde iki eksende yetersiz kalmaktadır. "
    "Birincisi şeffaflık eksikliğidir: kullanıcı neden belirli bir içeriği gördüğünü "
    "genellikle bilmez. İkincisi refah-körü tasarımdır: sıralama yalnızca etkileşimi "
    "optimize eder, kullanıcının o anki duygusal kırılganlığını hiç hesaba katmaz. "
    "Sonuç olarak, kullanıcı refahını sıralama kararının kendisine entegre eden, "
    "açıklanabilir ve içerik-duyarlı bir çözüm, piyasada hâlâ ele alınmamış gerçek bir "
    "problem olarak durmaktadır; bu proje tam olarak bu boşluğu hedef almaktadır."
)

altbaslik("2.2. Çözüm Fikri, Özgünlük ve Yerlilik")
govde(
    "Geliştirilen çözüm, iki katmanlı açıklanabilir bir sıralama motorudur. İlgi skoru, "
    "kullanıcının platformda beyan ettiği ilgi alanlarıyla (örneğin spor, teknoloji, "
    "gündem) bir gönderinin konusunun ne kadar örtüştüğünü ölçer; klasik bir "
    "kişiselleştirme sinyali gibi çalışır. Refah skoru ise tamamen farklı bir soruya "
    "cevap verir: kullanıcı şu anda olası bir \"spiral\" içinde mi? \"Spiral\", "
    "kullanıcının kısa sürede art arda, olumsuz ve yoğun duygulu içerikte uzun süre "
    "durup neredeyse hiç aktif tepki vermediği; doomscrolling olarak bilinen davranışa "
    "benzeyen bir davranışsal örüntüyü tanımlar (bkz. Bölüm 3.2). Bu örüntü, eğitilmiş "
    "bir sınıflandırıcı tarafından 0 ile 1 arasında bir \"spiral seviyesi\" olarak "
    "tahmin edilir; seviye yükseldikçe kullanıcının ilgi alanının içinde kalınarak yalnızca "
    "daha az tetikleyici içeriğe doğru bir kayma sağlanır, ilgi alanı hiç terk edilmez. "
    "İlgi skoru ile refah skoru birlikte gönderinin final skorunu oluşturur ve akış bu "
    "skora göre sıralanır."
)
govde(
    "Motorun akademik bir emsali vardır: \"Collective Well-Being aware "
    "Recommendation Systems (CWB-RS)\" kavramı, etkileşim optimizasyonu yerine uzun "
    "vadeli kümülatif refahı maksimize etmeyi önerir [7]; endüstri/politika düzeyinde de "
    "benzer \"insan-öncelikli\" sıralama ilkeleri savunulmaktadır [8]. Özgünlüğün temel "
    "kaynağı, kararların nasıl üretildiğidir. Birçok tavsiye/sıralama sistemi, kendi "
    "kararını kendisi de tam olarak açıklayamayan büyük ve karmaşık modellere (örneğin "
    "derin sinir ağlarına) dayanır; bu modeller kullanıcıya yalnızca \"bana güven\" "
    "diyebilir, \"neden\" sorusuna gerçek bir cevap veremez. Bizim sistemimizde ise her "
    "karar (spiral seviyesi, refah cezası, final skor) denetlenebilir, açıklanabilir "
    "modellerle (lojistik regresyon/SGD) üretilir ve şeffaflık panelinde kullanıcıya "
    "gösterilir."
)
govde(
    "Mevcut piyasa çözümleriyle karşılaştırıldığında (bkz. Bölüm 2.1) fark somuttur: "
    "Instagram'ın Take a Break'i ve TikTok'un Scheduled Breaks'i yalnızca geçen süreye "
    "bakan, içerikten ve kullanıcının anlık duygu durumundan tamamen bağımsız, isteğe "
    "bağlı hatırlatıcılardır; kullanıcı bunları tek dokunuşla görmezden gelebilir. Bizim "
    "çözümümüz müdahaleyi ayrı bir hatırlatma katmanına değil, sıralamanın kendisine "
    "gömer: karar geçen süreye değil içeriğin duygusal yoğunluğuna ve kullanıcının o anki "
    "davranışsal örüntüsüne dayanır, kullanıcıdan hiçbir ek adım beklemeden otomatik "
    "işler. Çözümün pazarda uygulanabilir olması, prototipin NSosyal'in mevcut arayüz "
    "mantığına bilinçli olarak sadık kalınarak tasarlanmış olmasından (bkz. Bölüm 3.3) ve "
    "önerilen düşük maliyetli sürdürülebilirlik modelinden (bkz. Bölüm 6.1) "
    "anlaşılabilir."
)
govde("Yerlilik bileşenleri somuttur:")
madde([
    "NSosyal'in kendisi T3 Vakfı ve Baykar Teknoloji tarafından geliştirilen yerli bir "
    "platformdur.",
    "Duygu analizi için kullanılan temel model Türkçeye özel eğitilmiştir ve bu model, "
    "bağımsız doğrulamada bulunan bir zayıflık (bkz. Bölüm 3.2) üzerine ekibimiz "
    "tarafından yeniden eğitilerek (fine-tuning) doğruluğu %69,8'den %94,3'e "
    "çıkarılmıştır; yani kullanılan yapay zekâ bileşeni yalnızca tüketilen değil, yerel "
    "olarak geliştirilen bir teknik varlıktır.",
    "NSosyal'in mevcut T3 AI bileşeni (otomatik yanıt, çok dilli yorumlama, filtreleme/"
    "moderasyon), duygu-analizi katmanımızın önerilen bir uzantısı olarak "
    "konumlandırılmıştır.",
])

# ============================================================ 3. TEKNOLOJİ KULLANIMI
baslik("TEKNOLOJİ KULLANIMI")

altbaslik("3.1. İzlenecek Yöntem, Altyapı ve Sürüm Kontrolü")
govde(
    "Backend Python/FastAPI ile geliştirilmiştir; duygu analizi için HuggingFace "
    "Transformers (savasy/bert-base-turkish-sentiment-cased tabanlı [15], ince ayarlı), "
    "sınıflandırıcılar için scikit-learn (LogisticRegression / SGDClassifier), veri "
    "işleme için NumPy/Pandas kullanılmıştır. Fine-tuning ve bağımsız doğrulama için "
    "winvoker/turkish-sentiment-analysis-dataset (HuggingFace Hub) [16] kullanılmıştır; "
    "spiral ve psikolojik izlenim sınıflandırıcıları için ise senaryo-bazlı sentetik veri "
    "üretimi tercih edilmiştir (bkz. Bölüm 3.2). Frontend, dwell-time takibi için "
    "tarayıcının yerel Intersection Observer API'sini kullanan sade bir JavaScript/HTML/"
    "CSS arayüzüdür. Haftalık öz-farkındalık raporu, oturum verisini okunabilir bir metne "
    "çeviren Google Gemini API'si (gemini-3.6-flash) ile üretilir; sağlayıcı seçimi "
    "bilinçlidir: Gemini'nin kart istemeyen gerçek bir ücretsiz katmanı, açıklanabilirlik "
    "ilkesiyle çelişmeyen, sağlayıcıdan bağımsız tasarlanmış bir prompt mimarisiyle "
    "kullanılmaktadır."
)
govde(
    "Sistemin teknik altyapısı, aşağıdaki uçtan uca veri akışıyla çalışır: kullanıcı "
    "arayüzdeki bir gönderiyle etkileşime girer; tarayıcı bu sinyali (durma süresi, "
    "tıklama, roket, yorum) bir REST API çağrısıyla backend'e iletir; backend gerekli "
    "modelleri (duygu, spiral, psikolojik izlenim) çalıştırıp güncellenmiş skorları ve "
    "bir doğal-dil açıklaması geri döner; arayüz bu yanıtla \"tespit edilen durum\" "
    "göstergesini, renk doygunluğunu ve şeffaflık panelini anlık olarak günceller. Bu "
    "akışı taşıyan başlıca uç noktalar:"
)
tablo(
    ["Uç nokta", "İşlev"],
    [
        ["GET /api/gonderiler", "İlgi + refah skoruna göre sıralanmış, sayfalanmış gönderi akışını döner"],
        ["POST /api/etkilesim", "Bir etkileşimi işler; spiral seviyesini ve psikolojik izlenimi günceller"],
        ["POST /api/dogrulama", "Kullanıcının onay cevabını modelin o anki tahminiyle karşılaştırır"],
        ["GET /api/psikolojik-ozet", "Oturumda gerçekten gözlemlenen kategori dağılımını döner"],
        ["GET /api/haftalik-rapor", "Oturum verisinden LLM ile gerçek zamanlı öz-farkındalık raporu üretir"],
        ["GET /api/terapist-raporu", "Aynı veriden, yorumsuz/ham bir uzman-veri özeti üretir"],
        ["POST /api/kisisel-mod", "Geçmişi Varsayılan/Kişiselleştirilmiş modelle yeniden skorlar"],
    ],
)
govde(
    "Proje sürüm kontrolü altında geliştirilmiş, düzenli ve anlamlı commit'lerle "
    "ilerletilmiştir; kaynak kod GitHub üzerinde açık bir depoda saklanmaktadır: "
    "https://github.com/samibahar/nsosyal. Depo, commit geçmişiyle geliştirme sürecinin "
    "tamamını (ilk prototip aşamasından güncel sürüme kadar) takip edilebilir kılar."
)

altbaslik("3.2. Model ve Veri Doğrulama")
govde(
    "Veri ön işleme iki farklı boru hattı izler. Duygu modeli tarafında, BERT'in kendi "
    "alt-sözcük (WordPiece) tokenizer'ı metni sayısal girdilere çevirir; fine-tuning için "
    "winvoker veri setinden dengeli (7.500 pozitif + 7.500 negatif) bir örneklem alınır ve "
    "dinamik dolgulama (dynamic padding) ile eğitim hızlandırılır. Davranışsal "
    "sınıflandırıcılar tarafında ise eğitim verisi, her kategori için gerçekçi davranışsal "
    "imzalar tanımlayan (örneğin \"sinirli\": negatif duygu + sık roket/yorum + kısa-orta "
    "dwell) senaryo üreteçleriyle sentetik olarak oluşturulur, ardından %10 kasıtlı etiket "
    "gürültüsü eklenir. Ham özellikler (dwell_saniye 0-15+ saniye ile duygu -1..1 ve ikili "
    "sinyaller 0/1) çok farklı ölçeklerde olduğundan, eğitim öncesi StandardScaler ile "
    "ölçeklenir; bu adım atlandığında psikolojik izlenim modelinin F1 makro değeri "
    "0,686'da kalırken, uygulandığında 0,704'e çıktığı bağımsız olarak ölçülmüştür."
)
govde(
    "Proje, vendor/ilk-varsayım iddialarını sorgulamadan kabul etmek yerine her bileşeni "
    "bağımsız olarak ölçme disiplinini benimser. Duygu modeli için: temel modelin kendi "
    "iddia ettiği doğruluk (%95,4) yerine, winvoker/turkish-sentiment-analysis-dataset "
    "üzerinden bağımsız 1000 örneklik bir test setiyle gerçek doğruluk ölçülmüş (%69,8, "
    "F1 0,778) ve alan kayması (domain shift) kaynaklı bu zayıflık, aynı veri setinin "
    "ayrık bir bölümüyle (veri sızıntısı önlenerek) ~15.000 örnekle yapılan hedefli bir "
    "ince ayarla düzeltilmiştir (%94,3 doğruluk, F1 0,964; bağımsız olarak yeniden "
    "ölçülmüştür). Spiral tespiti (olası doomscrolling örüntüsü, ikili sınıflandırma), "
    "senaryo-bazlı sentetik veriyle eğitilmiş bir lojistik regresyon modelidir "
    "(doğruluk 0,714, F1 0,748); altı açıklanabilir özellik (negatif dwell oranı, "
    "ortalama duygu, tıklama oranı vb.) kullanır ve öğrenilen katsayılar denetlenebilir "
    "durumdadır. Beş kategorili psikolojik izlenim sınıflandırıcısı (sakin/mutluluk/"
    "umut/sinirli/anksiyete), yukarıda anlatılan ölçekleme ile eğitilmiş bir "
    "SGDClassifier'dır (F1 makro 0,704). Aşırı öğrenmeyi (overfitting) önlemek için: "
    "tüm modellerde stratified train/test ayrımı ile fine-tuning ve bağımsız doğrulama "
    "için birbirinden tamamen ayrık veri bölümleri kullanılmıştır (yukarıdaki %10 etiket "
    "gürültüsü de aynı amaca, sınırların asla net olmadığını modellemeye, hizmet eder). "
    "Modelin gerçek "
    "performansını iddia etmek yerine sürekli ölçmek için, kullanıcıya ara sıra hafif "
    "bir onay sorusu (\"şu an gerçekte nasıl hissediyorsun?\") sorulur ve cevap modelin "
    "o anki tahminiyle karşılaştırılır; bu mekanizma psikolojideki Ecological Momentary "
    "Assessment (EMA) yönteminin bağımsız bir uygulamasıdır [1][2]; literatür bu tür "
    "tekrarlı öz-bildirim sorgularının kullanıcı yorgunluğuna (fatigue) yol açabileceğini "
    "de not eder [3], bu yüzden soru sıklığı bilinçli olarak düşük (her ~8 etkileşimde "
    "bir) tutulmuştur."
)

t_ozet = tablo(
    ["Bileşen", "Ölçülen sonuç", "Doğrulama yöntemi"],
    [
        ["Duygu modeli (BERT)", "%94,3 doğruluk (ince ayar sonrası, %69,8'den)",
         "Bağımsız 1000 örnek, vendor iddiasından ayrı ölçüldü"],
        ["Spiral tespiti", "Doğruluk 0,714 / F1 0,748", "Stratified train/test, senaryo verisi"],
        ["Psikolojik izlenim (5 kategori)", "F1 makro 0,704",
         "Ölçek düzeltmesi (StandardScaler) ile 0,686'dan iyileştirildi"],
        ["Kendi kendini doğrulama döngüsü", "Canlı test edildi, çalışır durumda",
         "Gerçek etkileşimde model tahmini ile kullanıcı onayı eşleşti"],
    ],
)

altbaslik("3.3. Kullanıcı Deneyimi (UI/UX) Tasarımı")
govde(
    "Kullanıcı akışı şu şekilde işler:"
)
madde([
    "Kullanıcı ana akışa girer; gönderiler ilgi ve refah skoruna göre sıralanmış olarak "
    "karşısına çıkar.",
    "Bir gönderiyle etkileşime girer (durur, tıklar, roket atar veya yorum yazar).",
    "Sistem bu etkileşimi anında işler; üstteki \"tespit edilen durum\" göstergesi ve "
    "akışın renk doygunluğu güncellenir.",
    "\"Neden bunu görüyorsun?\" butonuna basarak, o gönderiye özel şeffaflık panelini "
    "(ilgi skoru/refah cezası/final skoru) açabilir.",
    "Yaklaşık her 8 etkileşimde bir, kısa ve geçilebilir bir onay sorusu görünür.",
    "İstediği an \"Haftalık Rapor\" sayfasına geçip o oturumda gerçekten gözlemlenen "
    "örüntüleri, model doğrulama istatistiklerini ve gerçek zamanlı üretilen yapay zekâ "
    "yorumunu görebilir.",
])
govde(
    "Tasarım kararlarının gerekçesi şu şekildedir: arayüz, kullanıcıyı yeni bir "
    "paradigmayla karşılaştırmamak için NSosyal'in mevcut zaman akışı (timeline) "
    "mantığına bilinçli olarak sadık kalır; şeffaflık, akışın dışına taşan ayrı bir "
    "sayfa yerine her kartın içinde, isteğe bağlı açılan bir panel olarak sunulur, böylece "
    "istemeyen kullanıcı hiç rahatsız edilmez. Renk doygunluğu azaltma da aynı ilkeyle "
    "(kullanıcı isteğiyle açılıp kapatılabilir bir anahtarla) engelleyici değil, fark "
    "ettirici bir sinyal olacak şekilde tasarlanmıştır. Kullanılabilirlik testleri "
    "sırasında "
    "(gerçek kullanım denemeleriyle) birden fazla somut sorun tespit edilip düzeltilmiştir: "
    "aynı gönderiye art arda etkileşim (roket sonra kaydırıp uzaklaşma) önceden sinyali "
    "sessizce eziyordu, sinyal birleştirme mantığıyla giderildi; çok kısa/edilgen bir "
    "görünme (<0,6sn) tam bir gözlem gibi sayılıp barları saptırıyordu, etkileşim-ağırlıklı "
    "ortalamayla düzeltildi; ilk sürümde sayfalama sığ kalıyordu, gerçek sonsuz-kaydırma "
    "(infinite scroll) ile 150 örnek gönderiye genişletildi. Erişilebilirlik açısından: "
    "tüm renk kodlu göstergeler (spiral seviyesi, kategori barları) aynı zamanda metinsel "
    "etiket taşır (yalnızca renge dayanmaz), ve şeffaflık paneli literatürdeki \"az ama "
    "net\" ilkesine uygun kısa/öz tutulmuştur; aşırı detaylı açıklamaların kullanıcı "
    "güvenini azaltabileceği (\"algorithmic aversion\") bulgusuyla uyumludur [9]."
)

# ============================================================ 4. UYGULANABİLİRLİK
baslik("UYGULANABİLİRLİK")

altbaslik("4.1. Verimlilik ve Etkinlik")
govde(
    "Sistemin sıralama/tespit katmanı bilinçli olarak hafif ve açıklanabilir modeller "
    "(lojistik regresyon, SGD) üzerine kurulmuştur. Somut bir verimlilik farkı şudur: "
    "spiral sınıflandırıcısı yalnızca altı özellik ve bir sabit terimden, toplam yedi "
    "öğrenilmiş parametreden oluşur (bkz. Bölüm 3.2); her sıralama kararını bunun yerine "
    "bir büyük dil modeliyle almak, milyarlarca parametreli bir modeli her istek için "
    "çalıştırmayı gerektirirdi. Bu fark sayesinde sistem milisaniyeler mertebesinde "
    "tahmin üretir ve bir GPU'ya ihtiyaç duymadan standart sunucu donanımında "
    "çalışabilir. LLM çağrısı yalnızca kullanıcı başına haftada bir üretilen "
    "öz-farkındalık raporunda ve isteğe bağlı (talep üzerine, otomatik değil) uzman "
    "özetinde kullanılır; bu da toplam çağrı hacmini düşük tutar. Bu mimari tercih, "
    "platforma milyonlarca etkileşimde dahi düşük işletme maliyetiyle ölçeklenebilirlik "
    "sağlar; NSosyal'in ~700 bin'i aşkın kullanıcısına, mimarinin doğası gereği büyük "
    "bir altyapı yatırımı olmadan teorik olarak hizmet verebilecek bir temel sunar; bu "
    "iddia şu aşamada gerçek bir yük testiyle değil, mimari tasarımla desteklenmektedir."
)
govde(
    "Etkinlik, uydurma olmayan somut bir sayıyla ölçülebilir biçimde ifade edilebilir: "
    "refah cezası, spiral seviyesi ile içeriğin negatif duygu yoğunluğunun çarpımının "
    "0,8 katsayısıyla ölçeklenmesinden hesaplanır; yani en yüksek spiral seviyesinde "
    "(1,0) ve en olumsuz içerikte final skordan 0,8 puana kadar somut bir ceza "
    "uygulanır ve içerik sırada geriye düşer. Aynı mekanizmada kullanıcının ilgi alanı "
    "hiç terk edilmez; içerik elenmez, yalnızca yumuşatılır. Bu, sistemin \"etkileşimi "
    "düşürmeden refahı önceleme\" hedefinde etkin olduğunun somut bir kanıtıdır: "
    "kullanıcı istediği konuyu görmeye devam eder, yalnızca o konunun en tetikleyici "
    "örnekleri geri plana çekilir."
)

altbaslik("4.2. Hedef Kitle")
govde(
    "Bu çözümün nihai hedef kitlesi NSosyal kullanıcılarıyla sınırlı değildir: sorunun "
    "kökü, yani \"en yüksek etkileşim\" hedefiyle optimize edilen ve olumsuz/"
    "kutuplaştırıcı içeriği yapısal olarak ödüllendiren akış tasarımı (bkz. Bölüm 1.1 ve "
    "2.1), platform bağımsız ve küresel bir olgudur; bu tüketim örüntüsünden etkilenen "
    "herkes çözümün asıl adresidir. Bu kitlenin büyüklüğü soyut değildir: yarışma "
    "şartnamesinin de belirttiği gibi sosyal medya ekosistemi günümüzde milyarlarca "
    "kullanıcı tarafından aktif olarak kullanılmaktadır; NSosyal'in ~700 bin'i aşkın "
    "kullanıcısı bu küresel kitlenin somut ve erişilebilir bir alt kümesidir. Bu "
    "vizyonu somut ve test edilebilir bir noktadan "
    "hayata geçirmek için, NSosyal'in ~700 bin'i aşkın aktif kullanıcısı doğrudan/ilk "
    "uygulama kitlesi olarak seçilmiştir; özellikle gündem/haber içeriğiyle yoğun "
    "etkileşime giren, uzun oturumlar geçiren kullanıcı segmenti önceliklidir. Dolaylı "
    "olarak, sistemin ürettiği şeffaflık ve refah verileri içerik üreticileri (üretici "
    "paneli önerisi) ve platform yöneticileri için de değer taşır. Mimari tasarımı "
    "gereği herhangi bir sosyal medya platformuna uyarlanabilir olduğundan, NSosyal "
    "ötesine büyüme somut bir mimari tercihe dayanır; yalnızca soyut bir iddia değildir. "
    "Ürünün hedef kitleyle uyumu da varsayılmamış, test edilmiştir: arayüzün NSosyal'in "
    "mevcut kullanım mantığına kasıtlı olarak sadık kalınarak tasarlanması ve "
    "kullanılabilirlik testlerinde bu uyumun somut sorunlar üzerinden doğrulanıp "
    "düzeltilmiş olması (bkz. Bölüm 3.3), uyumun kanıtlandığının bir göstergesidir."
)

altbaslik("4.3. Teknolojik Yenilik ve Uygulanabilirlik")
govde(
    "Ürün, fikir düzeyinde kalmayıp uçtan uca çalışan bir prototiple desteklenmiştir: "
    "gerçek bir BERT modeli, eğitilmiş sınıflandırıcılar, çalışan bir backend/API, "
    "gerçek zamanlı bir arayüz ve canlı test edilmiş bir LLM entegrasyonu içerir. "
    "Teknolojik yenilik düzeyi, tek bir modelin doğruluğunda değil, iki ayrı davranış "
    "modelinin (spiral tespiti ve psikolojik izlenim) birbirinden bağımsız çalışıp yalnızca "
    "biri sıralamayı etkileyecek şekilde tasarlanmasında, ve kendi kendini doğrulama "
    "döngüsüyle modelin gerçek performansının sürekli ölçülmesinde yatar. Mimari, yeni "
    "sinyaller (örn. ek etkileşim türleri) veya yeni kategoriler eklemeye açık, "
    "ölçeklenebilir bir yapıya sahiptir; kişiselleştirme mekanizması (online öğrenme, "
    "bkz. Bölüm 6.2) bu ölçeklenebilirliğin somut bir kanıtıdır."
)
govde(
    "Prototip, bu vizyonun tamamını değil, temel mekanizmasını kanıtlar. Zaman ve kapsam "
    "kısıtları nedeniyle şu an eklenmeyen, ama tasarım aşamasında değerlendirilip somut "
    "birer gelecek iş maddesi olarak tanımlanan yönler şunlardır:"
)
madde([
    "Perspektif Köprüsü: kullanıcı bir konuda olumsuz veya çaresizlik hissi veren "
    "içerikte takılı kaldığında, aynı konunun çözüme dönük bir yönünü gösteren bir "
    "gönderiyi köprü olarak sunmak. Sistemin konu-nötr ilkesini siyasi taraf değil "
    "duygusal ton ekseninde koruyabildiği için mevcut mimariye uygun görünmektedir; ancak "
    "hangi sinyalin \"çözüme dönük\" içeriği belirleyeceği ayrı bir tasarım çalışması "
    "gerektirir.",
    "Kişiselleştirmenin üretim kalitesine taşınması (bkz. Bölüm 6.2): şu anki online "
    "öğrenme prototipi, düzenlileştirme (regularization) ve kullanıcı bazlı kalibrasyon "
    "eklenerek gerçek, çok kullanıcılı bir sisteme dönüştürülebilir.",
    "Gizlilik önceliğinin artırılması: analiz bileşenlerinin bir kısmının kullanıcı "
    "cihazında çalışması ve kişiselleştirmenin federe öğrenmeyle merkezi bir sunucu "
    "olmadan paylaşılması. Mevcut prototip şu an tamamen sunucu taraflı çalışmaktadır; "
    "bu ayrımın dürüstçe belirtilmesi gerekir.",
    "Dijital Refah Ağacı: haftalık/aylık metin raporunun yanına, refah eğiliminin zaman "
    "içindeki seyrini temsil eden görsel bir metafor eklemek. Mevcut metin ve çubuk "
    "grafik raporu literatürdeki \"az ama net\" ilkesine zaten uygun olduğundan, bu "
    "eklenti karmaşıklaştırma riskine karşı dikkatli tasarlanmalıdır.",
    "Dinamik Sürtünme: renk doygunluğu azaltmanın ötesinde, spiral seviyesi çok "
    "yükseldiğinde kaydırma deneyimine küçük bir duraklama veya onay adımı eklemek. Bu, "
    "mevcut \"engellemek değil fark ettirmek\" ilkesine göre daha agresif bir müdahaledir; "
    "kullanıcı özerkliğini kısıtlama riski nedeniyle ayrı bir kullanıcı araştırması "
    "gerektirdiğinden bu aşamada önceliklendirilmemiştir.",
])

# ============================================================ 5. YAYGIN ETKİ
baslik("YAYGIN ETKİ")

altbaslik("5.1. Toplumsal Fayda ve Erişim Potansiyeli")
govde(
    "Bu bölümde ele alınan zarar varsayımsal değildir; sağlık otoriteleri tarafından "
    "resmi olarak kayıt altına alınmıştır. ABD Sağlık Bakanlığı Baş Cerrahı'nın 2023 "
    "tarihli tavsiye raporu, 13-17 yaş grubundaki gençlerin %95'inin sosyal medya "
    "kullandığını, üçte birinin \"neredeyse sürekli\" kullandığını, ergenlerin %64'ünün "
    "nefret söylemi içeren içerikle karşılaştığını tespit etmiş; platform "
    "algoritmalarının kullanıcıyı saatlerce meşgul edecek şekilde tasarlandığını ve "
    "bunun uyku ile dikkat üzerinde ölçülebilir bir bedeli olduğunu vurgulayarak "
    "durumu gençler için \"derin bir zarar riski\" olarak tanımlamıştır [19]. Dünya "
    "Sağlık Örgütü Avrupa Bölge Ofisi'nin verileri aynı yönde ilerlemektedir: "
    "ergenlerde sorunlu sosyal medya kullanımı oranı 2018'den 2022'ye %7'den %11'e "
    "yükselmiş, kız çocuklarındaki oran (%13) erkek çocuklarına (%9) göre belirgin "
    "biçimde yüksek çıkmıştır [20]."
)
govde(
    "Sorunun platform içi kaynağına dair kanıt da dolaylıdan ibaret değildir. Bölüm "
    "2.1'de değinilen Facebook'un kendi iç araştırması, 2020 tarihli bir sunumda teen "
    "kız kullanıcıların %32'sinin, vücut imajı hakkında kötü hissettiklerinde "
    "Instagram'ın bu hissi daha da kötüleştirdiğini bildirdiğini; intihar düşüncesi "
    "bildiren gençler arasında bu düşünceyi doğrudan Instagram'a bağlayanların "
    "oranının İngiltere'de %13, ABD'de %6 olduğunu ortaya koymuştur [13]. Bu, "
    "dışarıdan yapılmış bir tahmin değil, platformun kendi verisiyle ve kendi "
    "araştırmacıları tarafından üretilmiş bir bulgudur; şirketin bu bulguları kamuya "
    "açıklamadığı ve küçümsediği de aynı belgelerde yer almaktadır."
)
govde(
    "Sorunun büyüklüğü Türkiye bağlamında da somuttur. TÜİK'in 2024 Hanehalkı "
    "Bilişim Teknolojileri Kullanım Araştırması'na göre Türkiye'de 16-74 yaş "
    "grubunun %88,8'i internet, %74,4'ü Instagram kullanmaktadır; bu yaklaşık 57,5 "
    "milyon sosyal medya kullanıcısına karşılık gelir [25]. Türkiye'deki gençlerin "
    "sosyal medya alışkanlıklarını inceleyen bir derleme, gençlerin yaklaşık "
    "yarısının günde dört saatini sosyal medyada geçirdiğini, yaklaşık %12'sinin ise "
    "günde yedi saatten fazla zaman harcadığını aktarmaktadır [24]. Bu ölçek, hem "
    "sorunun hem de önerilen çözümün erişim potansiyelinin NSosyal'in mevcut "
    "kullanıcı tabanının çok ötesine uzandığını somutlaştırır (bkz. Bölüm 4.2)."
)
govde(
    "Bu noktada asıl soru şudur: içeriği kaldırmayan, yalnızca sıralamasına dokunan "
    "bir müdahale gerçekten duygusal sonuçları değiştirebilir mi? Bu sorunun cevabı "
    "artık spekülatif değildir. 2024 ABD başkanlık seçim kampanyası sırasında "
    "yürütülen, 1.256 katılımcılı, 10 günlük önden kayıtlı bir saha deneyinde "
    "araştırmacılar, bir tarayıcı eklentisiyle X platformundaki akışı gerçek "
    "zamanlı olarak yeniden sıraladı; yalnızca karşıt partiye yönelik düşmanlık "
    "içeren gönderilerin akıştaki sıklığını artırıp azaltarak (içerik silmeden, "
    "yalnızca sırasını değiştirerek) katılımcıların karşıt partiye yönelik duygusal "
    "tutumlarını 100 puanlık bir ölçekte 2 puandan fazla kaydırmayı başardılar; bu "
    "etki büyüklüğü, ABD'de yaklaşık üç yıl boyunca biriken duygusal kutuplaşma "
    "değişimiyle kıyaslanabilir düzeydedir [21]. Bu çalışma, bizim sistemimizin "
    "dayandığı temel varsayımla doğrudan örtüşür: bir içeriği kaldırmadan, yalnızca "
    "akıştaki konumunu ve dolayısıyla maruziyet miktarını değiştirmek, ölçülebilir "
    "duygusal sonuçlar üretebilecek gerçek bir mekanizmadır. Refah katmanımız da tam "
    "olarak bunu yapar; farkı, hedefinin siyasi düşmanlık değil olumsuz ve "
    "tetikleyici duygusal yoğunluk olmasıdır."
)
govde(
    "Müdahalenin içeriği kaldırmak yerine yönlendirmeyi tercih etmesi de literatürde "
    "ayrıca desteklenir. Doğruluk odaklı hafif hatırlatmaların yanlış bilgi "
    "paylaşılma niyetini kontrol grubuna göre yaklaşık %10 oranında azalttığı, bu "
    "etkinin 20 deneyi ve 26.863 katılımcıyı kapsayan bir meta-analizde demografik "
    "gruplar arasında tutarlı biçimde gözlemlendiği raporlanmıştır [22]. Bu araştırma "
    "bizim doğrudan alanımız olan duygusal spiral tespiti değil yanlış bilgi "
    "paylaşımı üzerinedir; yine de ortak ilke aynıdır: kullanıcının özerkliğini "
    "ortadan kaldırmayan, dikkati yeniden yönlendiren hafif müdahaleler ölçülebilir "
    "davranışsal etkiler üretebilir."
)
govde(
    "Bu iyimser tabloyu dengelemek için karşıt bulgular da dürüstçe belirtilmelidir. "
    "X üzerinde 4.965 katılımcıyla yaklaşık 7 hafta süren büyük ölçekli bir saha "
    "deneyinde, kullanıcıların kronolojik ile algoritmik akış arasında geçiş yapması "
    "öznel refah üzerinde anlamlı bir fark yaratmamıştır; oysa aynı deney siyasi "
    "tutumlarda anlamlı kaymalar üretmiştir [26]. Benzer biçimde, Facebook'un kendi "
    "\"Snooze\" ve süre hatırlatma araçlarının kullanımı, bir kesitsel araştırmada "
    "depresyon, kaygı veya stres puanlarıyla anlamlı biçimde ilişkilendirilememiştir "
    "[27]. Bu iki bulgu rastgele değil, tutarlı bir örüntüye işaret eder: yalnızca "
    "akış mantığını kronolojiye çevirmek veya isteğe bağlı, içerikten bağımsız bir "
    "hatırlatıcı eklemek (bkz. Bölüm 2.1'deki mevcut çözümler), tek başına refahı "
    "ölçülebilir biçimde iyileştirmeye yetmemektedir. Tasarımımız bilinçli olarak bu "
    "iki yaklaşımdan da ayrılır: akışı kronolojiye çevirmez, kişiselleştirmeyi korur; "
    "isteğe bağlı ve görmezden gelinebilir bir hatırlatıcı da eklemez, müdahaleyi "
    "doğrudan sıralama kararının içine, yalnızca içeriğin duygusal yoğunluğuna göre "
    "hedeflenmiş biçimde gömer. Bu, yukarıdaki örneklerin ortak eksiğini, yani "
    "hedefsizliği, doğrudan adresleyen bir tasarım kararıdır."
)
govde(
    "Şeffaflık panelinin tasarım gerekçesi de yalnızca sezgiye değil, doğrudan "
    "etkilenen kullanıcılarla yapılan araştırmaya dayanır. Ruh sağlığı tanısı almış "
    "21 katılımcıyla yürütülen bir tasarım atölyesi çalışması, kullanıcıların kendi "
    "eylemleri ile platformun onlara gösterdiği sonuçlar arasında bir kopukluk "
    "yaşadığını; bu kopukluğu gidermek için önerdikleri çözümlerin tam olarak "
    "etkileşimi bağlamsallaştırmak ve kullanıcıya açık bir kontrol iadesi etmek "
    "üzerine kurulduğunu ortaya koymuştur [23]. Şeffaflık panelimiz (bkz. Bölüm 2.2 "
    "ve 3.3) bu iki talebi doğrudan karşılar: her karar açıklanabilir sunulur ve "
    "kullanıcı doygunluk azaltmayı istediği an kapatabilir."
)
govde(
    "Bu temel üzerine, toplumsal fayda üç düzeyde değerlendirilebilir. Bireysel "
    "düzeyde, haftalık öz-farkındalık raporu kullanıcıya kendi dijital alışkanlıkları "
    "hakkında olası örüntüleri (teşhis değil) gösterir ve isteğe bağlı olarak bir "
    "ruh sağlığı uzmanına götürülebilecek, yorumsuz bir ham veri özeti üretebilir. "
    "Platform düzeyinde, refah-farkında sıralama yaklaşımının akademik emsali [7][8] "
    "literatürde giderek daha fazla destek bulmaktadır. Toplumsal düzeyde ise "
    "konu-nötr ve reklam-manipülasyonuna kapalı tasarım ilkesi (bkz. Bölüm 2.2), "
    "yerli bir platformun kullanıcı refahını önceleyen bir marka değeri inşa "
    "etmesine katkı sağlayabilir. Erişim potansiyeli, sistemin NSosyal'in mevcut "
    "kullanıcı tabanının tamamına, herhangi bir ek donanım veya kullanıcı eğitimi "
    "gerektirmeden arka planda çalışan bir katman olarak doğrudan ulaşabilecek "
    "şekilde tasarlanmış olmasından kaynaklanır; Bölüm 4.2'de belirtildiği gibi bu "
    "erişim, mimarinin platform bağımsız taşınabilirliği sayesinde NSosyal'in "
    "ötesine de genişleyebilir."
)
govde(
    "Son olarak kapsam dürüstçe sınırlanmalıdır: yukarıdaki kaynaklar bizim "
    "spesifik sistemimizin gerçek kullanıcılarda ruh sağlığını iyileştirdiğini "
    "kanıtlamaz; bu tür bir kanıt ancak uzun soluklu, gerçek kullanıcı "
    "popülasyonunda yürütülecek bir çalışmayla üretilebilir ve şu an elimizde "
    "değildir. Elimizde olan, üç ayrı düzeyde yakınsayan dolaylı ama tutarlı bir "
    "kanıt zinciridir: sorunun gerçek ve büyük ölçekli olduğunu gösteren resmi ve "
    "akademik kaynaklar; aynı müdahale sınıfının (kaldırmadan yeniden sıralama, "
    "hafif yönlendirme) başka bağlamlarda ölçülebilir etkiler ürettiğini gösteren "
    "nedensel çalışmalar; ve bu iki bulguyu birleştiren, hedefli ve açıklanabilir "
    "bir uygulama olarak bizim sistemimiz. Bölüm 3.2'deki kendi kendini doğrulama "
    "döngüsü, bu zincirin son halkasını, yani gerçek etkinin zamanla ölçülebilmesini "
    "sağlayacak altyapıyı oluşturur."
)

# ============================================================ 6. SÜRDÜRÜLEBİLİRLİK
baslik("SÜRDÜRÜLEBİLİRLİK")

altbaslik("6.1. Ticarileştirme Potansiyeli ve İş Modeli")
govde(
    "NSosyal vakıf tabanlı ve reklamsız bir platform olduğundan klasik reklam "
    "gelirine dayalı bir model zorlanmamıştır; bunun yerine dört tamamlayıcı "
    "gelir/değer akışı önerilmektedir. Birincisi, düşük işletme maliyetidir: "
    "hafif/açıklanabilir modeller büyük dil modelleri kadar pahalı değildir, bu "
    "doğrudan bir maliyet avantajına dönüşür (bkz. Bölüm 4.1). İkincisi, "
    "anonim/toplu \"dijital refah eğilimleri\" verisiyle üniversite, TÜBİTAK veya "
    "ruh sağlığı sivil toplum kuruluşlarıyla araştırma ortaklığıdır (kişisel veri "
    "paylaşılmadan, yalnızca toplu istatistik düzeyinde); bu model varsayımsal "
    "değildir, Meta'nın Center for Open Science ile kurduğu ve Instagram verisini "
    "gizlilik koruyucu biçimde akademik araştırmacılara açtığı refah odaklı pilot "
    "programı, büyük platformların bu tür ortaklıkları gerçekten kurduğunu "
    "göstermektedir [28]. Üçüncüsü, \"dijital-refah-öncelikli platform\" marka "
    "değeri üzerinden kurumsal sponsorluktur. Dördüncüsü, mevcut roket sistemine "
    "benzer, gönüllülük esaslı opsiyonel bir destek/bağış mekanizmasıdır; bu "
    "modelin ölçeği küçümsenmemelidir: Wikimedia Foundation, benzer bir gönüllü "
    "bağış modeliyle 2023-2024 mali yılında 8 milyondan fazla bağışçıdan, "
    "ortalama 10,58 dolarlık katkılarla 168,21 milyon dolarlık bir gelir elde "
    "etmiştir [29]. Bu model, ürünün mevcut pazar şartlarında (vakıf destekli, "
    "reklamsız yerli platform bağlamında) üretilebilir olmasını gerçekçi bir "
    "temele oturtur."
)
govde(
    "Sektöre ve ülke ekonomisine katma değer potansiyeli de somuttur. Yarışma "
    "şartnamesinin kendisi, Türkiye'nin sosyal medya teknolojileri alanındaki "
    "bilgi birikimini artırarak küresel ölçekte rekabet edebilecek teknolojilerin "
    "geliştirilmesini ve bu projelerin ticarileşerek ülkenin teknoloji "
    "girişimciliği ekosistemine kazandırılmasını temel bir hedef olarak "
    "tanımlamaktadır (bkz. Bölüm 1.1). Bu proje, yerli bir platform üzerinde, "
    "yerli olarak ince ayarı yapılmış bir yapay zekâ bileşeniyle (bkz. Bölüm 2.2) "
    "çalıştığından, doğrudan bu hedefe katkı sunacak biçimde konumlanmıştır: "
    "yurt dışı kaynaklı, kapalı kutu duygu analizi/refah araçlarına bağımlılığı "
    "azaltan, yerli ve açıklanabilir bir alternatif ortaya koyar. İş ortaklıkları "
    "açısından da somut adaylar mevcuttur: üniversitelerin psikoloji ve bilgisayar "
    "mühendisliği bölümleri, TÜBİTAK destekli araştırma programları, ruh sağlığı "
    "alanında çalışan sivil toplum kuruluşları ve \"dijital refah\" odaklı "
    "kurumsal sponsorluk arayan markalar, bu projenin gerçekçi ilk işbirliği "
    "adaylarıdır."
)

altbaslik("6.2. Finansal, Teknik ve Sosyal Sürdürülebilirlik")
govde(
    "Finansal sürdürülebilirlik, düşük işletme maliyetli mimari tercihinden "
    "(büyük/pahalı modeller yerine açıklanabilir, hafif sınıflandırıcılar) "
    "doğrudan gelir; Bölüm 4.1'de gösterildiği gibi bu, milyonlarca etkileşimde "
    "dahi düşük maliyetle çalışabilmeyi ve dolayısıyla uzun vadeli "
    "işletilebilirliği mümkün kılar."
)
govde(
    "Teknik sürdürülebilirlik iki ayrı boyutta ele alınmalıdır: bakım ve "
    "ölçeklenme. Bakım tarafında, projenin kendi bulgusu somut bir uyarı "
    "niteliğindedir: Bölüm 3.2'de bağımsız doğrulamada ortaya çıkan alan kayması, "
    "bir modelin bir kez eğitilip sonsuza kadar güvenilir kalamayacağını "
    "göstermiştir; platformun içerik dağılımı zamanla değiştikçe modellerin "
    "periyodik olarak yeniden doğrulanması ve gerektiğinde yeniden ince ayar "
    "yapılması, tek seferlik değil sürekli bir bakım süreci olarak "
    "planlanmalıdır. Kendi kendini doğrulama döngüsü (bkz. Bölüm 3.2) bu bakım "
    "sürecine gereken erken uyarı sinyalini sağlar: model tahminleri ile "
    "kullanıcı onayları arasındaki uyuşma oranı zamanla düşerse, bu durum "
    "yeniden eğitim ihtiyacının doğrudan bir göstergesi olur. Ölçeklenme "
    "tarafında ise, Bölüm 4.1'de detaylandırılan hafif model mimarisi ve Bölüm "
    "4.3'te belirtilen yeni sinyal/kategori eklemeye açık yapı, teknik "
    "sürdürülebilirliğin diğer yarısını oluşturur."
)
govde(
    "Sosyal sürdürülebilirlik açısından, sistemin kişiselleştirme mekanizması "
    "(SGDClassifier'ın partial_fit özelliğiyle çevrimiçi öğrenme) örnek bir "
    "prototip olarak uygulanmış ve dürüstçe sınırları belirtilmiştir. Literatür, "
    "davranıştan duygu durumu çıkarımının meşru ama kişiye özgü kalibrasyon "
    "olmadan sınırlı kaldığını göstermektedir [6]; kişiselleştirme mekanizması "
    "tam olarak bu ihtiyaca yönelik bir ilk adımdır. Demo ortamında öğrenme adımı "
    "görünür olması için kasıtlı büyütülmüştür, gerçek üretimde çok daha küçük "
    "bir adım büyüklüğü ve düzenlileştirme/sınır (clipping) eklenmesi gerektiği "
    "açıkça not edilmiştir; bu, değişen kullanıcı ihtiyaçlarına zamanla uyum "
    "sağlayabilecek bir temel sunar. Modelin diğer sınırlılıkları (sentetik "
    "eğitim verisi, klinik doğrulama eksikliği) raporun her aşamasında dürüstçe "
    "belirtilmiş, kendi kendini doğrulama döngüsüyle sistemin zamanla gerçek "
    "kullanıcı geri bildirimiyle kalibre edilebilecek bir altyapı kurulmuştur; "
    "bu, kullanıcı güveninin uzun vadede korunmasını destekler."
)

# ============================================================ 7. PROJE TAKVİMİ
baslik("PROJE TAKVİMİ")

altbaslik("7.1. İş Paketleri ve Zamanlama")
govde(
    "Aşağıdaki takvim, yarışma takvimiyle (Teknik Rapor Teslimi: 24 Ağustos 2026; "
    "Mentörlük Süreci: 2-7 Eylül 2026; Final Sunumları: 14 Eylül 2026) çelişmeyecek "
    "şekilde planlanmıştır. Tamamlanan iş paketleri, bu raporun dayandığı çalışan "
    "prototipi kapsar; ilerleyen iş paketleri mentörlük ve final aşamasına "
    "hazırlığı hedefler."
)

t_takvim = tablo(
    ["İş Paketi", "Kapsam", "Dönem", "Durum"],
    [
        ["İP1: Mimari ve ilk prototip",
         "İki katmanlı sıralama motoru, spiral/psikolojik model tasarımı, sözlük-tabanlı ilk doğrulama",
         "Ağustos (1. hafta)", "Tamamlandı"],
        ["İP2: Gerçek model entegrasyonu",
         "BERT entegrasyonu, bağımsız doğrulama, ince ayar, sınıflandırıcı eğitimi",
         "Ağustos (2. hafta)", "Tamamlandı"],
        ["İP3: Backend/arayüz ve kullanıcı testi",
         "FastAPI backend, web arayüzü, kullanılabilirlik testleri ve düzeltmeler",
         "Ağustos (2-3. hafta)", "Tamamlandı"],
        ["İP4: Şeffaflık, rapor ve doğrulama döngüsü",
         "LLM destekli haftalık rapor, kendi kendini doğrulama döngüsü, kişiselleştirme",
         "Ağustos (3. hafta)", "Tamamlandı"],
        ["İP5: Teknik rapor ve sürüm kontrolü",
         "Teknik rapor yazımı, GitHub deposunun düzenlenmesi ve teslimi",
         "24 Ağustos 2026'ya kadar", "Tamamlandı"],
        ["İP6: Mentörlük geri bildirimlerinin uygulanması",
         "Mentör önerileri doğrultusunda model/mimari iyileştirmeleri",
         "2-7 Eylül 2026", "Planlandı"],
        ["İP7: Kişiye-özgü kalibrasyon ve genişletme",
         "Kişiselleştirme mekanizmasının düzenlileştirme ile güçlendirilmesi, gelecek-"
         "vizyonu özelliklerinin (Perspektif Köprüsü vb.) fizibilite çalışması",
         "Eylül (1-2. hafta)", "Planlandı"],
        ["İP8: Final sunum ve canlı demo hazırlığı",
         "Sunum dosyası, kullanıcı senaryoları, canlı demo akışının hazırlanması",
         "14 Eylül 2026'ya kadar", "Planlandı"],
    ],
)
govde(
    "Bu iş paketlerini somut kilometre taşlarına bağlamak gerekirse: Ağustos "
    "ayının ikinci haftası sonunda (İP1-İP2) duygu modelinin bağımsız "
    "doğrulaması ve ince ayarı tamamlanmış olacaktır; üçüncü hafta başında "
    "(İP3-İP4) kullanılabilirlik testleriyle düzeltilmiş, uçtan uca çalışan bir "
    "prototip ve kendi kendini doğrulama döngüsü hazır olacaktır; 24 Ağustos "
    "2026'da teknik rapor ve sürüm kontrollü kaynak kod KYS üzerinden teslim "
    "edilecektir; 2-7 Eylül 2026'da mentörlük geri bildirimleri uygulanmış "
    "olacaktır; 14 Eylül 2026'da final sunum dosyası ve canlı demo akışı hazır "
    "olacaktır. Bu tarihler yarışmanın resmi takvimiyle (24 Ağustos, 2-7 Eylül, "
    "14 Eylül 2026) doğrudan uyumludur ve yukarıdaki iş paketi tablosundaki "
    "dönemlerle birebir örtüşür."
)

# ============================================================ 8. TAKIM YAPISI
baslik("TAKIM YAPISI")

altbaslik("8.1. Takım Organizasyonu ve Roller")
govde(
    "Değerlendirme esasları gereği takım üyelerinin isim ve fotoğraf gibi kişisel "
    "bilgilerine bu bölümde yer verilmemiştir. Takım, yarışma kayıt şartına uygun "
    "olarak 2-5 kişi ve bir takım kaptanından oluşmaktadır; görev dağılımı aşağıda "
    "disiplin/rol bazında özetlenmiştir."
)

t_takim = tablo(
    ["Rol", "Katkı Alanı"],
    [
        ["Takım Kaptanı / Proje Yönetimi", "Genel koordinasyon, yarışma süreç takibi, teslim yönetimi"],
        ["Yazılım Geliştirme (Backend)", "FastAPI backend, API tasarımı, sürüm kontrolü"],
        ["Yapay Zekâ / Veri Bilimi", "Model eğitimi, bağımsız doğrulama, ince ayar, performans ölçümü"],
        ["Ürün / UI-UX Tasarımı", "Arayüz tasarımı, kullanıcı akışları, kullanılabilirlik testleri"],
        ["Araştırma ve Dokümantasyon", "Literatür taraması, teknik rapor, kaynakça"],
    ],
)
notv(
    "Not: Tablo, projede yürütülen disiplin bazlı katkı alanlarını göstermektedir; "
    "ekip büyüklüğü ve isimlendirme yarışma kayıt sistemindeki resmi bilgilerle "
    "tutarlıdır."
)

sayfa_sonu()

# ============================================================ KAYNAKÇA
baslik("KAYNAKÇA")

kaynaklar = [
    "[1] Investigating Best Practices for Ecological Momentary Assessment, JMIR, "
    "26:e50275, 2024. Erişim: https://www.jmir.org/2024/1/e50275",
    "[2] Measuring Criterion Validity of Microinteraction Ecological Momentary "
    "Assessment (Micro-EMA), PMC7991987. Erişim: "
    "https://www.ncbi.nlm.nih.gov/pmc/articles/PMC7991987/",
    "[3] Modeling Behaviour to Predict User State: Self-Reports as Ground Truth, "
    "arXiv:2007.14461. Erişim: https://arxiv.org/pdf/2007.14461",
    "[4] Are active and passive social media use related to mental health, wellbeing, "
    "and social support outcomes? A meta-analysis, Journal of Computer-Mediated "
    "Communication, 29(1), zmad055, 2024. Erişim: "
    "https://academic.oup.com/jcmc/article/29/1/zmad055/7595758",
    "[5] Beyond the Scroll: Intolerance of Uncertainty and Doomscrolling, "
    "ScienceDirect, 2024. Erişim: "
    "https://www.sciencedirect.com/science/article/abs/pii/S0191886924003799",
    "[6] Passive Sensing for Mental Health Monitoring Using Machine Learning with "
    "Wearables and Smartphones: A Scoping Review, JMIR, 2025. Erişim: "
    "https://www.jmir.org/2025/1/e77066",
    "[7] Challenging Social Media Threats using Collective Well-being Aware "
    "Recommendation Algorithms and Multi-objective Optimization, arXiv:2102.04211. "
    "Erişim: https://arxiv.org/pdf/2102.04211",
    "[8] Better Feeds: Algorithms That Put People First, Georgetown University "
    "Knight-Georgetown Institute, Mart 2025, Erişim Tarihi: 19.08.2026, Erişim: "
    "https://kgi.georgetown.edu/wp-content/uploads/2025/02/Better-Feeds_-Algorithms-That-Put-People-First.pdf",
    "[9] Explainable recommendation: when design meets trust calibration, World Wide "
    "Web (Springer), 2021. Erişim: "
    "https://link.springer.com/article/10.1007/s11280-021-00916-0",
    "[10] Robertson, C.E., Pröllochs, N., Schwarzenegger, K., Pärnamets, P., Van "
    "Bavel, J.J., Feuerriegel, S., (2023) Negativity drives online news consumption, "
    "Nature Human Behaviour, 7, 812-822. Erişim: "
    "https://www.nature.com/articles/s41562-023-01538-4",
    "[11] Negative online news articles are shared more to social media, Scientific "
    "Reports (Nature), 2024. Erişim: https://www.nature.com/articles/s41598-024-71263-z",
    "[12] Negativity Spreads More than Positivity on Twitter After Both Positive and "
    "Negative Political Situations, PMC9383030. Erişim: "
    "https://www.ncbi.nlm.nih.gov/pmc/articles/PMC9383030/",
    "[13] The Facebook Files, The Wall Street Journal, 2021, Erişim Tarihi: "
    "19.08.2026, Erişim: https://www.wsj.com/articles/the-facebook-files-11631713039",
    "[14] Confidential Facebook document reveals leverage over vulnerable teens, "
    "The Australian, 2017 (haber kaynağı; ABD Kongresi'ne sunulan iç belgelerle "
    "ilişkilendirilmiştir).",
    "[15] savasy/bert-base-turkish-sentiment-cased, HuggingFace model kartı, Erişim "
    "Tarihi: 19.08.2026, Erişim: https://huggingface.co/savasy/bert-base-turkish-sentiment-cased",
    "[16] winvoker/turkish-sentiment-analysis-dataset, HuggingFace veri seti, Erişim "
    "Tarihi: 19.08.2026, Erişim: https://huggingface.co/datasets/winvoker/turkish-sentiment-analysis-dataset",
    "[17] Instagram tests 'Take a Break' reminders on an opt-in basis, TechCrunch, "
    "10.11.2021, Erişim Tarihi: 19.08.2026, Erişim: "
    "https://techcrunch.com/2021/11/10/instagram-tests-take-a-break-reminders-on-an-opt-in-basis",
    "[18] Helping users manage their screen time, TikTok Newsroom, 2022, Erişim Tarihi: "
    "19.08.2026, Erişim: https://newsroom.tiktok.com/en-us/helping-users-manage-their-screen-time",
    "[19] Social Media and Youth Mental Health: The U.S. Surgeon General's Advisory, "
    "ABD Sağlık Bakanlığı (HHS), 2023, Erişim Tarihi: 20.08.2026, Erişim: "
    "https://www.hhs.gov/sites/default/files/sg-youth-mental-health-social-media-advisory.pdf",
    "[20] Teens, screens and mental health, WHO Avrupa Bölge Ofisi, 25.09.2024, Erişim "
    "Tarihi: 20.08.2026, Erişim: "
    "https://www.who.int/europe/news/item/25-09-2024-teens--screens-and-mental-health",
    "[21] Piccardi, T., Saveski, M., Jia, C., Hancock, J.T., Tsai, J.L., Bernstein, M.S., "
    "Reranking partisan animosity in algorithmic social media feeds alters affective "
    "polarization, Science, DOI: 10.1126/science.adu5584. Erişim: "
    "https://www.science.org/doi/10.1126/science.adu5584",
    "[22] Pennycook, G., Rand, D.G., (2022) Nudging Social Media toward Accuracy, The "
    "ANNALS of the American Academy of Political and Social Science. Erişim: "
    "https://journals.sagepub.com/doi/10.1177/00027162221092342",
    "[23] Milton, A., Runningen, D., Terveen, L., Kaur, H., Chancellor, S., Unraveling "
    "Entangled Feeds: Rethinking Social Media Design to Enhance User Well-being, ACM CHI "
    "Conference on Human Factors in Computing Systems 2026. Erişim: "
    "https://arxiv.org/abs/2602.15745",
    "[24] Öztürk, S., Gençlerde Sosyal Medya Kullanımı ve Ruh Sağlığına Etkileri, SD "
    "(Sağlık Düşüncesi ve Tıp Kültürü) Dergisi, Sayı 66, 2024, s. 50-53, Erişim Tarihi: "
    "20.08.2026, Erişim: https://sdplatform.com/genclerde-sosyal-medya-kullanimi-ve-ruh-sagligina-etkileri/",
    "[25] Hanehalkı Bilişim Teknolojileri (BT) Kullanım Araştırması 2024, Türkiye "
    "İstatistik Kurumu (TÜİK), Erişim Tarihi: 20.08.2026, Erişim: "
    "https://data.tuik.gov.tr/Bulten/Index?p=Hanehalki-Bilisim-Teknolojileri-(BT)-Kullanim-Arastirmasi-2025-53925",
    "[26] The political effects of X's feed algorithm, Nature, 2026, PMC13061628. Erişim: "
    "https://pmc.ncbi.nlm.nih.gov/articles/PMC13061628/",
    "[27] Predicting Psychological Symptoms When Facebook's Digital Well-being Features "
    "Are Used: Cross-sectional Survey Study, PMC9468917. Erişim: "
    "https://www.ncbi.nlm.nih.gov/pmc/articles/PMC9468917/",
    "[28] Instagram Data Access Pilot for Well-being Research, Meta ve Center for Open "
    "Science (COS) ortaklığı, Erişim Tarihi: 20.08.2026, Erişim: "
    "https://www.cos.io/meta",
    "[29] 2023-2024 Annual Report, Wikimedia Foundation, Erişim Tarihi: 20.08.2026, "
    "Erişim: https://wikimediafoundation.org/annualreports/2023-2024-annual-report/",
]
for k in kaynaklar:
    p = doc.add_paragraph(k)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

doc.save(DOSYA)
print(f"Olusturuldu: {DOSYA}")
